#!/usr/bin/env python3
"""
Convert TECHNICAL_PAPER.md to TECHNICAL_PAPER.docx
"""
import re
import os
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Installing required packages...")
    import subprocess
    subprocess.check_call([
        'python', '-m', 'pip', 'install', 'python-docx', '--quiet'
    ])
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

def read_markdown(file_path):
    """Read markdown file."""
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()

def parse_markdown_to_docx(md_content, output_path):
    """Parse markdown and create Word document."""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    lines = md_content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Skip empty lines
        if not line.strip():
            i += 1
            continue
        
        # Title (# at start)
        if line.startswith('# '):
            title = line[2:].strip()
            p = doc.add_paragraph(title, style='Heading 1')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_format = p.paragraph_format
            p_format.space_before = Pt(12)
            p_format.space_after = Pt(12)
            i += 1
            
        # Heading 2 (## at start)
        elif line.startswith('## '):
            heading = line[3:].strip()
            p = doc.add_paragraph(heading, style='Heading 2')
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            
        # Heading 3 (### at start)
        elif line.startswith('### '):
            heading = line[4:].strip()
            p = doc.add_paragraph(heading, style='Heading 3')
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            i += 1
            
        # Heading 4 (#### at start)
        elif line.startswith('#### '):
            heading = line[5:].strip()
            p = doc.add_paragraph(heading, style='Heading 4')
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
            i += 1
            
        # Horizontal rule (---, ***, ___)
        elif re.match(r'^[\-\*_]{3,}$', line.strip()):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(12)
            i += 1
            
        # Code blocks (```)
        elif line.strip().startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph(code_text)
            p.style = 'Normal'
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
            p_format = p.paragraph_format
            p_format.left_indent = Inches(0.5)
            p_format.space_before = Pt(6)
            p_format.space_after = Pt(6)
            
        # Tables (| |)
        elif '|' in line:
            # Collect table rows
            table_rows = [line]
            i += 1
            while i < len(lines) and '|' in lines[i]:
                table_rows.append(lines[i])
                i += 1
            
            # Parse table
            rows = [r.strip('| ').split('|') for r in table_rows if r.strip()]
            if len(rows) > 0:
                # Create table
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = 'Table Grid'
                
                for row_idx, row_data in enumerate(rows):
                    cells = table.rows[row_idx].cells
                    for col_idx, cell_text in enumerate(row_data):
                        cells[col_idx].text = cell_text.strip()
                        # Make header row bold
                        if row_idx == 0:
                            for paragraph in cells[col_idx].paragraphs:
                                for run in paragraph.runs:
                                    run.font.bold = True
            
        # Bullet lists (- or *)
        elif line.strip().startswith(('- ', '* ')):
            bullet_text = line.strip()[2:].strip()
            p = doc.add_paragraph(bullet_text, style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            i += 1
            
        # Numbered lists (1. 2. etc)
        elif re.match(r'^\d+\.\s', line.strip()):
            match = re.match(r'^(\d+)\.\s(.+)$', line.strip())
            if match:
                num_text = match.group(2).strip()
                p = doc.add_paragraph(num_text, style='List Number')
                p.paragraph_format.space_after = Pt(3)
            i += 1
            
        # Regular paragraphs with formatting
        else:
            # Handle bold, italic, inline code
            paragraph_text = line.strip()
            if paragraph_text:
                p = doc.add_paragraph()
                
                # Parse inline formatting
                # Match bold **, italic *, code `
                pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|[^`*]+)'
                parts = re.findall(pattern, paragraph_text)
                
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        # Bold
                        run = p.add_run(part[2:-2])
                        run.font.bold = True
                    elif part.startswith('*') and part.endswith('*'):
                        # Italic
                        run = p.add_run(part[1:-1])
                        run.font.italic = True
                    elif part.startswith('`') and part.endswith('`'):
                        # Inline code
                        run = p.add_run(part[1:-1])
                        run.font.name = 'Courier New'
                        run.font.size = Pt(10)
                    else:
                        # Regular text
                        run = p.add_run(part)
                
                p.paragraph_format.space_after = Pt(6)
            
            i += 1
    
    # Save document
    doc.save(output_path)
    print(f"✓ Successfully created: {output_path}")

if __name__ == '__main__':
    input_file = 'TECHNICAL_PAPER.md'
    output_file = 'TECHNICAL_PAPER.docx'
    
    if not os.path.exists(input_file):
        print(f"Error: {input_file} not found")
        exit(1)
    
    print(f"Converting {input_file} to {output_file}...")
    parse_markdown_to_docx(read_markdown(input_file), output_file)
